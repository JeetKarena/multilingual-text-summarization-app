# BackEnd/app/models/document_parser.py
import os
import fitz  # PyMuPDF
import docx
from typing import Dict, Any

def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict containing extracted text and metadata
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        try:
            # Try to read as plain text
            return extract_text_from_txt(file_path)
        except:
            raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from a PDF file."""
    try:
        document = fitz.open(file_path)
        text = ""
        metadata = {
            "page_count": len(document),
            "title": document.metadata.get("title", ""),
            "author": document.metadata.get("author", ""),
            "subject": document.metadata.get("subject", ""),
            "keywords": document.metadata.get("keywords", "")
        }
        
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            text += page.get_text()
        
        document.close()
        return {
            "content": text,
            "metadata": metadata
        }
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract document properties
        core_properties = doc.core_properties
        metadata = {
            "title": core_properties.title if hasattr(core_properties, 'title') else "",
            "author": core_properties.author if hasattr(core_properties, 'author') else "",
            "subject": core_properties.subject if hasattr(core_properties, 'subject') else "",
            "paragraph_count": len(doc.paragraphs)
        }
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            "content": text,
            "metadata": metadata
        }
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        metadata = {
            "file_size": os.path.getsize(file_path),
            "line_count": content.count("\n") + 1
        }
        
        return {
            "content": content,
            "metadata": metadata
        }
    except Exception as e:
        raise Exception(f"Error extracting text from TXT: {str(e)}")

# BackEnd/app/models/summarization.py
import asyncio
import logging
from typing import Optional, Dict, Any, List
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.llms import HuggingFacePipeline
from app.core.config import settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Cache for loaded models
_models_cache = {}

async def generate_summary(
    text: str,
    language: str = "en",
    max_length: int = 150,
    min_length: int = 40
) -> str:
    """
    Generate a summary using the transformers pipeline directly.
    
    Args:
        text: Text to summarize
        language: Language code
        max_length: Maximum length of summary
        min_length: Minimum length of summary
        
    Returns:
        Generated summary
    """
    # Get the appropriate model for language
    model_key = f"mbart-{language}"
    
    # Initialize model if not cached
    if model_key not in _models_cache:
        await _initialize_model(model_key, language)
    
    # Get the pipeline from cache
    summarizer = _models_cache[model_key]["pipeline"]
    
    # Process text in chunks if too long
    if len(text) > 10000:
        return await _process_long_text(text, language, max_length, min_length)
    
    # Run summarization in executor to avoid blocking
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(
        None,
        lambda: summarizer(
            text,
            max_length=max_length,
            min_length=min_length,
            do_sample=False,
            truncation=True,
        )[0]['summary_text']
    )
    
    return result

async def generate_summary_with_langchain(
    text: str,
    language: str = "en",
    max_length: int = 150,
    min_length: int = 40,
    prompt_template: str = "Summarize the following text: {text}"
) -> str:
    """
    Generate a summary using LangChain integration.
    
    Args:
        text: Text to summarize
        language: Language code
        max_length: Maximum length of summary
        min_length: Minimum length of summary
        prompt_template: Custom prompt template for summarization
        
    Returns:
        Generated summary
    """
    # Get the appropriate model for language
    model_key = f"mbart-{language}"
    
    # Initialize model if not cached
    if model_key not in _models_cache:
        await _initialize_model(model_key, language)
    
    # Get the LangChain pipeline from cache
    langchain_pipeline = _models_cache[model_key]["langchain"]
    
    # Process text in chunks if too long
    if len(text) > 10000:
        return await _process_long_text(text, language, max_length, min_length)
    
    # Create a custom prompt if needed
    if prompt_template != "Summarize the following text: {text}":
        prompt = PromptTemplate(input_variables=["text"], template=prompt_template)
        langchain_pipeline = LLMChain(llm=_models_cache[model_key]["hf_pipeline"], prompt=prompt)
    
    # Run summarization in executor to avoid blocking
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(
        None,
        lambda: langchain_pipeline.run(text)
    )
    
    return result

async def detect_language(text: str) -> str:
    """
    Detect language of text using first few paragraphs.
    
    This is a simplified implementation. For production,
    consider using a dedicated language detection library.
    
    Args:
        text: Text to detect language
        
    Returns:
        Language code (ISO 639-1)
    """
    # Get a sample of text (first 1000 characters)
    sample = text[:1000]
    
    # Common language markers
    language_markers = {
        "the": "en", "and": "en", "for": "en",
        "le": "fr", "la": "fr", "et": "fr",
        "der": "de", "die": "de", "das": "de",
        "el": "es", "la": "es", "los": "es",
        "il": "it", "la": "it", "per": "it",
        "的": "zh", "是": "zh", "了": "zh",
        "は": "ja", "を": "ja", "が": "ja",
        "и": "ru", "в": "ru", "на": "ru"
    }
    
    # Count marker occurrences (very simplified)
    counts = {}
    words = sample.lower().split()
    for word in words:
        if word in language_markers:
            lang = language_markers[word]
            counts[lang] = counts.get(lang, 0) + 1
    
    # Return the language with most markers
    if counts:
        return max(counts, key=counts.get)
    
    # Default to English
    return "en"

async def _initialize_model(model_key: str, language: str):
    """
    Initialize and cache model for a specific language.
    """
    logger.info(f"Loading model for language: {language}")
    
    # Use the default multilingual model
    model_name = settings.DEFAULT_MODEL
    
    # Load tokenizer and model
    loop = asyncio.get_event_loop()
    tokenizer, model = await loop.run_in_executor(
        None,
        lambda: (
            AutoTokenizer.from_pretrained(model_name),
            AutoModelForSeq2SeqLM.from_pretrained(model_name)
        )
    )
    
    # Set up language if needed
    if language != "en":
        tokenizer.src_lang = language
    
    # Create the pipeline
    summarization_pipeline = pipeline(
        "summarization",
        model=model,
        tokenizer=tokenizer,
        framework="pt"
    )
    
    # Create LangChain components
    hf_pipeline = HuggingFacePipeline(pipeline=summarization_pipeline)
    prompt_template = PromptTemplate(
        input_variables=["text"],
        template="Summarize the following text: {text}"
    )
    langchain_pipeline = LLMChain(llm=hf_pipeline, prompt=prompt_template)
    
    # Cache everything
    _models_cache[model_key] = {
        "tokenizer": tokenizer,
        "model": model,
        "pipeline": summarization_pipeline,
        "hf_pipeline": hf_pipeline,
        "langchain": langchain_pipeline
    }
    
    logger.info(f"Model for {language} loaded successfully")

async def _process_long_text(
    text: str,
    language: str = "en",
    max_length: int = 150,
    min_length: int = 40
) -> str:
    """
    Process long text by splitting into chunks and summarizing each.
    
    Args:
        text: Long text to summarize
        language: Language code
        max_length: Maximum length of final summary
        min_length: Minimum length of final summary
        
    Returns:
        Combined summary
    """
    # Intelligent text splitting (by paragraphs)
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = ""
    
    # Create chunks of roughly 4000 chars (to stay under model limits)
    for para in paragraphs:
        if len(current_chunk + para) < 4000:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = para + "\n\n"
    
    # Add the last chunk if not empty
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # Summarize each chunk with smaller max_length
    summaries = []
    chunk_max_length = max(50, max_length // len(chunks))
    chunk_min_length = max(20, min_length // len(chunks))
    
    for i, chunk in enumerate(chunks):
        logger.info(f"Processing chunk {i+1}/{len(chunks)}")
        summary = await generate_summary(
            chunk,
            language=language,
            max_length=chunk_max_length,
            min_length=chunk_min_length
        )
        summaries.append(summary)
    
    # Combine summaries
    combined_text = " ".join(summaries)
    
    # If combined summaries still too long, summarize again
    if len(combined_text) > 4000:
        return await generate_summary(
            combined_text,
            language=language,
            max_length=max_length,
            min_length=min_length
        )
    
    return combined_text