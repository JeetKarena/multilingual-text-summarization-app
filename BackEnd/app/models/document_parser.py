# BackEnd/app/models/document_parser.py
import os
import fitz  # PyMuPDF
import docx
from typing import Dict, Any

def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict containing extracted text and metadata
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        try:
            # Try to read as plain text
            return extract_text_from_txt(file_path)
        except:
            raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from a PDF file."""
    try:
        document = fitz.open(file_path)
        text = ""
        metadata = {
            "page_count": len(document),
            "title": document.metadata.get("title", ""),
            "author": document.metadata.get("author", ""),
            "subject": document.metadata.get("subject", ""),
            "keywords": document.metadata.get("keywords", "")
        }
        
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            text += page.get_text()
        
        document.close()
        return {
            "content": text,
            "metadata": metadata
        }
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract document properties
        core_properties = doc.core_properties
        metadata = {
            "title": core_properties.title if hasattr(core_properties, 'title') else "",
            "author": core_properties.author if hasattr(core_properties, 'author') else "",
            "subject": core_properties.subject if hasattr(core_properties, 'subject') else "",
            "paragraph_count": len(doc.paragraphs)
        }
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            "content": text,
            "metadata": metadata
        }
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        metadata = {
            "file_size": os.path.getsize(file_path),
            "line_count": content.count("\n") + 1
        }
        
        return {
            "content": content,
            "metadata": metadata
        }
    except Exception as e:
        # Fix the missing quotation mark
        raise Exception(f"Error extracting text from TXT: {str(e)}")